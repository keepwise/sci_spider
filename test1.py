from docx.enum.text import WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx import Document
import tkinter
from tkinter import ttk
import threading
import tkinter.filedialog
import tkinter.messagebox

import numpy as np

import pandas as pd
#from docx.enum.text import WD_ALIGN_PARAGRAPH



'''''''''
@file
这个文件用来读取委托书书中的项目名称，委托人等信息，并写入excel文档

'''''''''

document = Document(r"C:\Users\wxs\Desktop\公司\查新\委托书\0304\大型垃圾焚烧发电厂施工关键技术研究委托合同.docx")
#document = Document(r"C:\1.docx")

tables = document.tables

table = tables[0]
weituo_data = []


for i in range(0, len(table.rows)):

    for j in range(0,len(table.rows[i].cells)):


        cellCont = str(table.rows[i].cells[j].text).strip()

        if i==0 and j==2:  weituo_data.append(cellCont)  #项目名称

        if i==2 and j ==2: weituo_data.append(cellCont)   #单位名称

        if i == 3 and j == 2: weituo_data.append(cellCont)  # 通信地址

        if i == 4 and j == 5: weituo_data.append(cellCont)  # 电子信箱

        if i == 5 and j == 2: weituo_data.append(cellCont)  # 负责人

        if i == 5 and j == 5: weituo_data.append(cellCont)  # 负责人手机

        if i == 6 and j == 2: weituo_data.append(cellCont)  # 联系人

        if i == 6 and j == 5: weituo_data.append(cellCont)  # 联系人手机

        if i == 7 and j == 2: weituo_data.append(cellCont)  # 开票信息



#print(weituo_data)

df = pd.DataFrame(np.array([weituo_data]),index=None, columns=['项目名称','委托人','地址','邮箱','负责人','负责人电话','联系人','联系人电话','开票信息'])

df.to_excel(r"C:\Users\wxs\Desktop\公司\查新\委托书\0304\1.xls")