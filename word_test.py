from docx import Document
import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time

document = Document()

document.styles["Normal"].font.name = "Times New Roman"
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
document.styles["Normal"].paragraph_format.space_after = docx.shared.Pt(1)

document.styles["Default Paragraph Font"].font.name = "Times New Roman"
document.styles['Default Paragraph Font']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

#header = document.sections[0].header
#header.text = "论文收录引用检索证明报告         2720"
#print(header.style.name)

print(document.sections[0].start_type)

p = document.add_paragraph("")
print(p.style.name)
print(p.style.type)
print(p.style.font.name)
p.add_run("附件一：").bold = True


p = document.add_paragraph("")

run = p.add_run("美国《科学引文索引》（SCI-EXPANDED）收录情况")
print(run.style.name)
print(run.style.font.name)
run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
run.bold = True

#设置引文段落格式
style = document.styles.add_style("yinwen", WD_STYLE_TYPE.PARAGRAPH)
style.paragraph_format.left_indent  = docx.shared.Cm(1)
style.font.name = "Times New Roman"
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

#设置标题样式
heading_style = document.styles.add_style("sci_heading", WD_STYLE_TYPE.CHARACTER)
heading_style.font.bold = True
heading_style.font.underline = True
heading_style.font.name = "Times New Roman"
heading_style.font.size =  docx.shared.Pt(14)
heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

style = document.styles.add_style("indent", WD_STYLE_TYPE.PARAGRAPH)
style.paragraph_format.left_indent = docx.shared.Cm(0.5)
style.paragraph_format.space_after = docx.shared.Pt(1)
style.font.name = "Times New Roman"
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = docx.shared.Pt(12)


p = document.add_paragraph()
p.add_run("检索结果", style="sci_heading")

p = document.add_paragraph("    本次检索根据委托人")
p.add_run("  ").underline = True
p.add_run("所提供的论文目录及其检索要求，通过对上面的数据库进行检索，检索结果如下：")

document.add_paragraph("1.	美国《科学引文索引》(SCI-EXPANDED，网络版)收录   篇，在《Web of Science 核心合集：引文索引》中累计被引用  次（其中他人引用  次，自引  次。注：关于他引和自引的区分，本证明所采用的方法是：文献被除第一作者及合作者以外其他人的引用为他引）。",style="indent")

document.add_paragraph("")
document.add_paragraph("（检索结果详见附件）")
document.add_paragraph("特此证明！")

document.add_paragraph("")
p = document.add_paragraph("        检索人：")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = document.add_paragraph("查证单位：教育部科技查新工作站（L27）")
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

p = document.add_paragraph("北京理工大学查新检索咨询中心")
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


p = document.add_paragraph(time.strftime("%Y{y}%m{m}%d{d}").format(y='年',m='月',d='日'))
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


document.save(r"c:\users\wangxiaoshan\desktop\wxs_py\rpt_test.docx")

print("success")
