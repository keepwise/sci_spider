import re
import urllib
import time
from datetime import datetime
import queue
from  lxml import etree
import urllib.robotparser
import urllib.parse
import urllib.request
import common
from docx import Document
import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import tkinter
from tkinter import ttk
import threading
import tkinter.filedialog
import tkinter.messagebox
import pymysql
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time


headers = {
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Encoding": "gzip, deflate",
    "Accept-Language": "zh-CN,zh;q=0.8,zh-TW;q=0.7,zh-HK;q=0.5,en-US;q=0.3,en;q=0.2",
    "Cache-Control": "max-age=0",
    "Connection": "keep-alive",
    "Upgrade-Insecure-Requests": "1",
    "content-type": "text/html;charset=UTF-8",
    "User-Agent": "Mozilla/5.0 (Windows NT 6.1; WOW64; rv:63.0) Gecko/20100101 Firefox/63.0"
}

url_crawled_num = 0  #爬取的原文数量
num_retries = 2   #爬取链接时尝试的次数
delay = 5   #请求间延迟的时间
proxy = None
wroten_original_num = 0  #已经写入的被引文献数量
document = Document()
document.styles["Normal"].font.name = "Times New Roman"
document.styles["Normal"].font.size = docx.shared.Pt(12)
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
document.styles["Normal"].paragraph_format.space_after = docx.shared.Pt(1)

document.styles["Default Paragraph Font"].font.name = "Times New Roman"
document.styles['Default Paragraph Font']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

original_papers_lst = []
cur_original_paper_no = 0
jcr_paragraph_position = 0
wos_cited_papers = 0  #SCI原文中，引用次数不为0的论文数量
processed_url_num = 0 #处理的链接数量

DEBUG = False

def get_papers_queue(seed_url):

    global url_crawled_num, num_retries, delay, proxy,headers


    throttle = common.Throttle(delay)

    url_queue = queue.deque()
    html = common.download(url=seed_url,proxy=None,num_retries=num_retries,headers=headers)
    html_emt = etree.HTML(html)
    page_title = html_emt.xpath("//title[contains(text(),'Citing Articles') or contains(text(),'施引文献')]")
    if page_title is None:
        return
    else:
        url_list = html_emt.xpath("//a[contains(@class,'snowplow-full-record')]/@href")
        for url in url_list:
            url_queue.append(url)
        next_page = html_emt.xpath("//a[contains(@class,'snowplow-navigation-nextpage-bottom')]/@href")

        try:
            if len(next_page) != 0:
                next_page = common.normalize(seed_url,next_page[0])
                throttle.wait(next_page)
                url_queue.extend(get_papers_queue(seed_url=next_page))
        except Exception as e:
            print("next_page: %s" % next_page)
            print(str(e))
            tkinter.messagebox.showinfo("错误","获取paper URL错误:%s %s" %(next_page, str(e)))
        return url_queue

def get_paper_record(url,database):
    global url_crawled_num, num_retries, delay, proxy, headers

    paper = {}
    html = common.download(url=url, proxy=None, num_retries=num_retries,headers=headers)
    html = html.encode(encoding='utf-8')
    html = html.decode("utf-8")

    #file = open(r"C:\Users\wangxiaoshan\Desktop\wxs_py\test.html", "w", encoding="utf-8")
    #file.write(html)
    html_emt = etree.HTML(html)
    try:

        title = html_emt.xpath("//div[@class='title']/value/text()")[0]
        #authors = html_emt.xpath( "//div[@class='l-content']//p[@class='FR_field']/span[contains(text(),'By:') or contains(text(),'作者:')]/following-sibling::a/text()|//div[@class='l-content']//p[@class='FR_field']/span[@id='more_authors_authors_txt_label']/a/text()")
        authors = html_emt.xpath("//div[@class='l-content']//p[@class='FR_field']/span[text()='By:' or text()='作者:']/following-sibling::a/text()|//div[@class='l-content']//p[@class='FR_field']/span[@id='more_authors_authors_txt_label']/a/text()")
        fullNames = ""
        print(authors)

        for author in authors:
            #fullName = html_emt.xpath('//p[@class="FR_field"]/span[contains(text(),"By:") or contains(text(),"作者:")]/following-sibling::a[text()="%s"]/following-sibling::text()|//div[@class="l-content"]//p[@class="FR_field"]/span[@id="more_authors_authors_txt_label"]/a[text()="%s"]/following-sibling::text()' % (author, author))[0]
            fullName = html_emt.xpath( '//p[@class="FR_field"]/span[text()="By:" or text()="作者:"]/following-sibling::a[text()="%s"]/following-sibling::text()|//div[@class="l-content"]//p[@class="FR_field"]/span[@id="more_authors_authors_txt_label"]/a[text()="%s"]/following-sibling::text()' % (author, author))[0]
            if fullName.find(";") != -1:

                fullNames = fullNames + author + fullName
            else:
                fullNames = fullNames + author + fullName + ";"

        print("FullNames:" + fullNames)

        SourceTitle = html_emt.xpath("//div[contains(@class,'block-record-info-source')]/p/span/value/text() | //div[contains(@class,'block-record-info-source')]/a/span/value/text() ")[0]

        SourceList = html_emt.xpath("//div[contains(@class,'block-record-info-source')]/p[@class='FR_field']/node()/text() | //div[contains(@class,'block-record-info-source')]/p[@class='FR_field']/text()")

        source = SourceTitle + " "

        i = 0
        for ss in SourceList:
            if i % 2:
                source += str(ss).replace("\n", "") + " "
            else:
                source += str(ss).replace("\n", "")
            i = i + 1

        citing_url = html_emt.xpath("//a[@class='snowplow-citation-network-times-cited-count-link']/@href")
        if len(citing_url):
            citing_url = citing_url[0]
        else:
            citing_url = ""

        reprint_author_lst = html_emt.xpath("//div[@class='block-record-info']//span[contains(text(),'Reprint Address')]/following-sibling::text()| //div[@class='block-record-info']//span[contains(text(),'通讯作者地址')]/following-sibling::text()")
        reprint_author_lst = filter(lambda x: "\n" not in str(x), reprint_author_lst)

        print(reprint_author_lst)
        reprint_addr_lst = html_emt.xpath("//table[@class='FR_table_noborders']/tr/td[2]/text()")

        rept_addr_author_lst = []
        rept_addr_author_lst = zip(reprint_author_lst,reprint_addr_lst)
        paper['reprint_author'] = ""

        for rpt in rept_addr_author_lst:
            if len(rpt)>1:
                paper['reprint_author'] += str(rpt[0]) + ", "
                paper['reprint_author'] += str(rpt[1]) + "\n"
            else:
                print("通讯作者错误: %s" % rept_addr_author_lst)
        print("reprint author: %s " % paper['reprint_author'])

        wos_cited_num = html_emt.xpath("//a[@class='snowplow-citation-network-times-cited-count-link']/span/text()")
        if len(wos_cited_num):
            wos_cited_num = wos_cited_num[0]
        else:
            wos_cited_num = "0"
        wos_no = html_emt.xpath("//input[@name='recordID']/@value")
        if len(wos_no):
            wos_no = wos_no[0]
        else:
            wos_no = ""
        data = html_emt.xpath("//div[@class='block-record-info']//span[contains(text(),'ISSN:')]/following-sibling::value/text()")
        if len(data):
            if len(data)==2:
                issn = data[0]
                eissn = data[1]
            if len(data)==1:
                issn = data[0]
                eissn = ""
        else:
            issn = ""
            eissn = ""

        paper['title'] = title
        paper['author'] = fullNames
        paper['source'] = source
        paper['citing_url'] = citing_url
        paper['wos_cited_num'] = wos_cited_num
        paper['wos_no'] = wos_no
        paper['issn']  = issn
        paper['eissn'] = eissn
        #自引次数
        paper['ziyin'] = 0
        #他引次数
        paper['tayin'] = 0
        paper['shoulu'] = database
    except Exception as e:
            print("获取文章详细信息失败,url: %s " % url)
            print("tite: %s" %  title)
            print("author: %s " % fullNames)
            print("source: %s" % source)
            print("错误：%s" % str(e))
            tkinter.messagebox.showinfo("错误","获取文章详细信息失败 %s" % url)
    return paper


def write_shoulu(document):

    global  wroten_original_num,original_papers_lst

    # 写入文件头
    p = document.add_paragraph("")
    p.add_run("附件一：").bold = True
    p = document.add_paragraph("")
    #run = p.add_run("美国《科学引文索引》（SCI-EXPANDED）收录情况")
    run = p.add_run("收录详细情况")
    run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
    run.bold = True

    style = document.styles.add_style("label",WD_STYLE_TYPE.CHARACTER)
    style.font.bold = True
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


    i = 0
    for paper in original_papers_lst:
        i += 1
        p = document.add_paragraph("")
        p.add_run(str(i)+". Title: ",style="label")
        p.add_run(paper['title'])

        p = document.add_paragraph("",style="indent")
        p.add_run("Author(s): ",style="label")
        p.add_run(paper['author'])

        p = document.add_paragraph("", style="indent")
        p.add_run("Source: ", style="label")
        p.add_run(paper['source'])

        if paper.get("wos_cited_num",'not') != 'not':
            p = document.add_paragraph("", style="indent")
            p.add_run("Web of Science Core Collection引用次数: ", style="label")
            p.add_run(paper['wos_cited_num'])

        if paper.get('wos_no','not') != 'not':
            p = document.add_paragraph("", style="indent")
            p.add_run("入藏号: ", style="label")
            p.add_run(paper['wos_no'])

        if paper.get('accession number','not') !='not':
            p = document.add_paragraph("", style="indent")
            p.add_run("Accession number: ", style="label")
            p.add_run(str(paper['accession number']))

        p = document.add_paragraph("", style="indent")
        p.add_run("通讯作者: ", style="label")
        p.add_run(str(paper['reprint_author']).replace("nan",""))

        p = document.add_paragraph("", style="indent")
        p.add_run("作者地址: ", style="label")
        p.add_run(str(paper['address']).replace("nan", ""))

        p = document.add_paragraph("", style="indent")
        p.add_run("ISSN: ", style="label")
        p.add_run(str(paper['issn']).replace("nan",""))

        if paper.get("eissn",'not') != 'not':
            p = document.add_paragraph("", style="indent")
            p.add_run("eISSN: ", style="label")
            p.add_run(paper['eissn'])
        if paper.get("hc","not")!="not":
            hc = str(paper['hc']).replace("nan","")
            if len(hc)>1:
                p = document.add_paragraph("", style="indent")
                p.add_run("是否高被引: ", style="label")
                p.add_run(paper['hc'])

        p = document.add_paragraph("", style="indent")
        p.add_run("收录情况: ", style="label")
        p.add_run(paper['shoulu'])

        document.save(str(gui.path_input.get()).strip()+"\\"+str(gui.bianhao_input.get()).strip()+"_baogao.docx")

def write_word(record, record_type,cite_total=0, cur_cite=0):
    #record_type注明是原文还是引文

    global document, wroten_original_num, original_papers_lst, cur_original_paper_no

    if record_type == "original":
        if wroten_original_num == 0:

            #写入文件头
            p = document.add_paragraph("")
            p.add_run("附件一：").bold = True
            p = document.add_paragraph("")
            run = p.add_run("美国《Web of Science 核心合集：引文索引》（网络版）引用情况")
            run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
            run.bold = True
            # 设置引文段落格式
            style = document.styles.add_style("yinwen", WD_STYLE_TYPE.PARAGRAPH)
            style.paragraph_format.left_indent = docx.shared.Cm(1)
            style.paragraph_format.space_after = docx.shared.Pt(1)

            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 设置标题样式
            heading_style = document.styles.add_style("sci_heading", WD_STYLE_TYPE.CHARACTER)
            heading_style.font.bold = True
            heading_style.font.underline = True
            heading_style.font.name = "Times New Roman"
            heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        wroten_original_num += 1
        p = document.add_paragraph("")
        p.add_run(str(wroten_original_num) + ". ")
        p.add_run("被引文献", style="sci_heading")
        p.add_run(":")

        p = document.add_paragraph("")
        p.add_run("Title: ").bold = True
        p.add_run(record['title'])

        p = document.add_paragraph("")
        p.add_run("Author(s): ").bold = True
        p.add_run(record['author'])

        p = document.add_paragraph("")
        p.add_run("Source: ").bold = True
        p.add_run(record['source'])

        p = document.add_paragraph("")
        p.add_run("WOS: ").bold = True
        p.add_run(record['wos_no'])

        p = document.add_paragraph("", style="yinwen")
        yinyong_paragraph_position = len(document.paragraphs)-1

        p.add_run("引用文献", style="sci_heading")
        p.add_run(": (自引").bold = True
        p.add_run("    ", style="sci_heading")
        p.add_run("篇  ").bold = True
        p.add_run("他引").bold = True
        p.add_run("    ", style="sci_heading")
        p.add_run("篇  )").bold = True

        original_papers_lst[cur_original_paper_no-1]['yinyong_paragraph_position'] = yinyong_paragraph_position
    if record_type == "citation":
        ziyin_tayin(record,cite_total,cur_cite)

    #print(str(gui.directory_input.get()).strip())
    document.save(str(gui.path_input.get()).strip() + "\\" + str(gui.bianhao_input.get()).strip()+"_yinyong.docx")

def ziyin_tayin(citation,cite_total=0, cur_cite=0):
    '''区分自引他引，并写入word文档'''
    global original_papers_lst,cur_original_paper_no

    original_authors = original_papers_lst[cur_original_paper_no-1]['author']
    original_authors_lst = original_authors.split(";")
    original_authors_lst = [str(x).strip() for x in original_authors_lst]
    #生成类似下述列表 [0]="Wang, XY (Wang, Xinyu)"   [1]="Fu, MY (Fu, Mengyin)" [2] = "Ma, HB (Ma, Hongbin)"

    citation_authors = citation['author']
    original_paper = original_papers_lst[cur_original_paper_no-1]

    p = document.add_paragraph("", style="yinwen")
    p.add_run("Record %d of %d" % (cur_cite, cite_total)).bold = True
    p = document.add_paragraph("", style="yinwen")
    p.add_run("Title: ").bold = True
    p.add_run(citation['title'])
    p = document.add_paragraph("", style="yinwen")
    p.add_run("Author(s): ").bold = True

    for full_author in original_authors_lst:
        # 如果[0]="Wang, XY (Wang, Xinyu)"   [1]="Fu, MY (Fu, Mengyin)" [2] = "Ma, HB (Ma, Hongbin)"匹配
        full_author_pattern = full_author.replace("-","[-\s]?")
        full_author_pattern = full_author_pattern.replace(", ","[,\s]*")
        full_author_pattern = full_author_pattern.replace(" ","[,\s]*")
        full_author_pattern = full_author_pattern.replace("(","\(")
        full_author_pattern = full_author_pattern.replace(")", "\)")

        if len(full_author)>1 and re.search(full_author_pattern,citation_authors,re.M|re.I) is not None:
            matchobj = re.search(full_author_pattern,citation_authors,re.M|re.I)
            full_author_matched = matchobj.group()
            author_len = len(full_author_matched)
            author_start = str(citation_authors).find(full_author_matched)
            p.add_run(citation['author'][0:author_start])
            p.add_run(citation['author'][author_start:(author_start+author_len)]).font.highlight_color = WD_COLOR_INDEX.RED
            p.add_run(citation['author'][(author_start + author_len):])

            original_papers_lst[cur_original_paper_no-1]['ziyin'] += 1
            break
    else:
        #如果[0]="Wang, XY (Wang, Xinyu)"   [1]="Fu, MY (Fu, Mengyin)" [2] = "Ma, HB (Ma, Hongbin)"不匹配
        # 就用 (Wang, Xinyu) ， (Fu, Mengyin)，(Ma, Hongbin)尝试匹配，如果匹配 就算自引一次
        for full_author in original_authors_lst:
            if str(full_author).find("(") != -1:
                full_name = str(full_author).split("(")[1]
                #full_name = (Wang, Xinyu)
                full_name = "(" + full_name.strip()
            else:
                full_name = str(full_author).strip()

            full_name = full_name.replace(", ","[,\s]*")
            full_name = full_name.replace("-","[-\s]*")
            full_name = full_name.replace(" ","[,\s]*")


            if len(full_name)>1 and re.search(full_name,citation_authors,re.M| re.I) is not None:

                full_name = re.search(full_name,citation_authors,re.M| re.I).group()
                author_len = len(full_name)
                author_start = str(citation_authors).find(full_name)
                p.add_run(citation['author'][0:author_start])
                p.add_run(citation['author'][author_start:(author_start + author_len)]).font.highlight_color = WD_COLOR_INDEX.RED
                p.add_run(citation['author'][(author_start + author_len):])

                original_papers_lst[cur_original_paper_no-1]['ziyin'] += 1
                break
            else:
                #如果full_name中含有",",如 Wang, Xinyu
                if full_name.find(",") != -1:
                    full_name = full_name.replace(",", "")
                    #对 Wang Xinyu进行检索
                    if str(citation_authors).find(full_name) != -1:
                        author_len = len(full_name)
                        author_start = str(citation_authors).find(full_name)
                        p.add_run(citation['author'][0:author_start])
                        p.add_run(citation['author'][author_start:(author_start + author_len)]).font.highlight_color = WD_COLOR_INDEX.RED
                        p.add_run(citation['author'][(author_start + author_len):])

                        original_papers_lst[cur_original_paper_no-1]['ziyin'] += 1
                        break
        else:
            # 如果[0]="Wang, XY (Wang, Xinyu)"   [1]="Fu, MY (Fu, Mengyin)" [2] = "Ma, HB (Ma, Hongbin)"不匹配
            # 如果 (Wang, Xinyu) ， (Fu, Mengyin)，(Ma, Hongbin)也不匹配
            #使用 Wang, XY,    Fu, MY, Ma, HB 进行匹配，两次匹配以上可以确认一次自引
            basic_matched_num = 0
            for full_author in original_authors_lst:
                if str(full_author).find("(") != -1:
                    basic_name = str(full_author).split("(")[0]
                    # 这里有问题要注意比如 Li, J 匹配上了Li, JF实际是不匹配的，应该测试Li, J (是否匹配
                    basic_name = basic_name.strip() + " ("
                    if str(citation_authors).find(basic_name) != -1:
                        author_len = len(basic_name)
                        author_start = str(citation_authors).find(basic_name)

                        p.add_run(citation['author'][0:author_start])
                        p.add_run(citation['author'][author_start:(author_start + author_len)]).font.highlight_color = WD_COLOR_INDEX.GREEN
                        p.add_run(citation['author'][(author_start + author_len):])

                        basic_matched_num += 1

                        if basic_matched_num>1:
                            original_papers_lst[cur_original_paper_no-1]['ziyin'] += 1
                            break
            else:
                p.add_run(citation['author'])

    p = document.add_paragraph("", style="yinwen")
    p.add_run("Source: ").bold = True
    p.add_run(citation['source'])

    p = document.add_paragraph("", style="yinwen")
    p.add_run("WOS: ").bold = True
    p.add_run(citation['wos_no'])

    original_papers_lst[cur_original_paper_no-1]['tayin'] = cite_total - original_papers_lst[cur_original_paper_no-1]['ziyin']
    yinyong_paragraph_position = original_papers_lst[cur_original_paper_no-1]['yinyong_paragraph_position']

    p = document.paragraphs[yinyong_paragraph_position]
    p = p.clear()
    p.add_run("引用文献", style="sci_heading")
    p.add_run(": (自引").bold = True
    p.add_run(" %d " % original_papers_lst[cur_original_paper_no-1]['ziyin'], style="sci_heading")
    p.add_run("篇  ").bold = True
    p.add_run("他引").bold = True
    p.add_run(" %d " % original_papers_lst[cur_original_paper_no-1]['tayin'], style="sci_heading")
    p.add_run("篇  )").bold = True

def get_wos_originals(path):
    global  wos_cited_papers
    csv_path = str(gui.wos_file_path.get()).encode(encoding="utf-8")
    csv_path = csv_path.decode("utf-8")
    f = open(csv_path,encoding="utf-8")
    papers_df = pd.read_csv(f,sep="\t",index_col=False)
    f.close()
    original_total = papers_df.shape[0]
    papers_lst = []

    if DEBUG == True:
        print(papers_df.head())
    i = 0
    while i<papers_df.shape[0]:
        paper = {}
        paper['title'] = papers_df['TI'][i]

        au = str(papers_df['AU'][i]).split(";")
        af = str(papers_df['AF'][i]).split(";")

        j = 0
        paper['author'] = ""
        for author in au:
            paper['author'] += str(author)+" ("+str(af[j])+");"
            j += 1

        paper['full_author'] = papers_df['AF'][i]
        so =  str(papers_df['SO'][i])
        vl =  str(papers_df['VL'][i])
        bp =  str(papers_df['BP'][i])
        ep = str(papers_df['EP'][i])
        py = str(papers_df['PY'][i])
        paper['source'] = so + " 卷:" + vl + " 页:" + bp + "-" + ep + " 出版年:" + py
        paper['wos_cited_num'] = str(papers_df['TC'][i])
        if paper['wos_cited_num'] != '0':
            wos_cited_papers += 1
        paper['ziyin'] = 0
        paper['tayin'] = 0
        paper['wos_no'] = str(papers_df['UT'][i])
        paper['issn'] = str(papers_df['SN'][i])
        paper['eissn'] = str(papers_df['EI'][i])

        paper['reprint_author'] = str(papers_df['RP'][i])
        paper['address'] = str(papers_df['C1'][i])
        paper['shoulu'] = "SCIE"

        #高被引
        if papers_df.get("HC") is not None:
            paper['hc'] = str(papers_df['HC'][i])
        if DEBUG == True:
            print(paper)

        i += 1
        papers_lst.append(paper)

    return papers_lst

def write_yinyong(paper,num,SID):
    global url_crawled_num, num_retries, delay, proxy, headers, processed_url_num,cur_original_paper_no
    cur_original_paper_no = num
    wos = paper['wos_no']
    time_now = datetime.now()
    cur_year = time_now.year
    url="http://apps.webofknowledge.com/WOS_GeneralSearch.do?fieldCount=1&action=search&product=WOS&search_mode=GeneralSearch&max_field_count=25&max_field_notice="+\
         "%E6%B3%A8%E6%84%8F%3A+%E6%97%A0%E6%B3%95%E6%B7%BB%E5%8A%A0%E5%8F%A6%E4%B8%80%E5%AD%97%E6%AE%B5%E3%80%82&input_invalid_notice="+\
         "%E6%A3%80%E7%B4%A2%E9%94%99%E8%AF%AF%3A+%E8%AF%B7%E8%BE%93%E5%85%A5%E6%A3%80%E7%B4%A2%E8%AF%8D%E3%80%82&exp_notice="+\
         "%E6%A3%80%E7%B4%A2%E9%94%99%E8%AF%AF%3A+%E4%B8%93%E5%88%A9%E6%A3%80%E7%B4%A2%E8%AF%8D%E5%8F%AF%E5%9C%A8%E5%A4%9A"+\
         "%E4%B8%AA%E5%AE%B6%E6%97%8F%E4%B8%AD%E6%89%BE%E5%88%B0+%28&input_invalid_notice_limits=+%3Cbr%2F%3E%E6%B3%A8%3A+%E6%BB%9A%E5%8A%A8%E6%A1%86%E4%B8%AD%E6%98%BE%E7"+\
         "%A4%BA%E7%9A%84%E5%AD%97%E6%AE%B5%E5%BF%85%E9%A1%BB%E8%87%B3%E5%B0%91%E4%B8%8E%E4%B8%80%E4%B8%AA%E5%85%B6%E4%BB%96%E6%A3%80%E7%B4%A2%E5%AD%97%E6%AE%B5%E7%9B%B8%E7"+\
         "%BB%84%E9%85%8D%E3%80%82&sa_params=WOS%7C%7C"+SID+"%7Chttp%3A%2F%2Fapps.webofknowledge.com%7C%27&formUpdated=true&value%28input1%29="+paper['wos_no']+"&value"+\
         "%28select1%29=UT&x=46&y=19&value%28hidInput1%29=&limitStatus=expanded&ss_lemmatization=On&ss_spellchecking=Suggest&SinceLastVisit_UTC=&SinceLastVisit_DATE=&range"+\
         "=ALL&period=Year+Range&startYear=1900&endYear="+str(cur_year)+"&editions=SCI&update_back2search_link_param=yes&ssStatus=display%3Anone&ss_showsuggestions=ON&"+\
        "ss_numDefaultGeneralSearchFields=1&ss_query_language=&rs_sort_by=PY.D%3BLD.D%3BSO.A%3BVL.D%3BPG.A%3BAU.A&SID="+SID

    throttle = common.Throttle(delay)

    url_queue = queue.deque()
    html = common.download(url=url, proxy=None, num_retries=num_retries, headers=headers)
    html = html.encode(encoding='utf-8')
    html = html.decode("utf-8")

    if DEBUG :
        f = open(r"C:\Users\wangxiaoshan\Desktop\wxs_py\tt.html","w",encoding="utf-8")
        f.write(html)
        f.close()
    html_emt = etree.HTML(html)

    try:
        citing_url = html_emt.xpath("//a[@class='snowplow-times-cited-link']/@href")
        if len(citing_url)<1:
            return False
        else:
            citing_url = citing_url[0]

        seed_url = "http://apps.webofknowledge.com"
        citing_url = common.normalize(seed_url, citing_url)

        citing_papers_queue = get_papers_queue(citing_url)

        if citing_papers_queue is not None:
            write_word(paper, record_type='original')
            cite_total = len(citing_papers_queue)
            cur_cite = 1
            while citing_papers_queue:
                citation_url = citing_papers_queue.popleft()
                citation_url = common.normalize(seed_url, citation_url)
                if DEBUG == True:
                    print(citation_url)
                throttle.wait(citation_url)
                citation = {}
                citation = get_paper_record(citation_url, "SCIE")
                processed_url_num += 1
                gui.processing_info.insert(0, "正在处理 %d: %s" % (processed_url_num, citation_url))
                write_word(citation, record_type='citation', cite_total=cite_total, cur_cite=cur_cite)
                cur_cite += 1
            return True
        else:
            return False
    except Exception as e:
        raise
        return  False

def get_wos_sid():
    global headers
    url = "http://webofknowledge.com/?DestApp=WOS&editions=SCI"
    request = urllib.request.Request(url=url, headers=headers)
    opener = urllib.request.build_opener()
    SID = ""
    try:
        response = opener.open(request)
        html = response.read().decode("UTF-8")
        #html = html.decode("utf-8")

        # file = open(r"C:\Users\wangxiaoshan\Desktop\wxs_py\test.html", "w", encoding="utf-8")
        # file.write(html)

        html_emt = etree.HTML(html)
        SID = html_emt.xpath("//input[@name='SID']/@value")[0]
    except Exception as e:
        tkinter.messagebox.showinfo("错误","获取SID错误")
        return None

    return  SID

def scrape_sci(seed_url):

    global original_papers_lst, wos_cited_papers
    try:
        # 如果WOS文件不为空
        if (len(gui.wos_file_path.get())>1):

            original_papers_lst = get_wos_originals(path=gui.wos_file_path.get())
            if gui.yinyong_opt.get() == True and wos_cited_papers != 0:
                SID = get_wos_sid()
                i = 1  # 引用报告中文献序号
                j = 1 #当期要处理的报告原文序号
                for paper in original_papers_lst:
                    if paper['wos_cited_num'] != '0':
                       if write_yinyong(paper, j,SID) == True:
                           i += 1
                    gui.progress_value.set((j / wos_cited_papers) * 100)
                    j += 1

        shoulu_document = Document()
        shoulu_document.styles["Normal"].font.name = "Times New Roman"
        shoulu_document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        shoulu_document.styles["Normal"].paragraph_format.space_after = docx.shared.Pt(1)

        shoulu_document.styles["Default Paragraph Font"].font.name = "Times New Roman"
        shoulu_document.styles['Default Paragraph Font']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


        ei_file = gui.ei_file_path.get()
        if(len(ei_file.strip())>1):
            ei_shoulu()

        write_report(shoulu_document)
        report_overview(shoulu_document)
        write_shoulu(shoulu_document)

        if gui.jcr_opt.get() == True:
            scrape_jcr(shoulu_document)
        if gui.fenqu_opt.get() == True:
            scrape_fenqu(shoulu_document)
        gui.processing_info.insert(0, "顺利完成")

        tkinter.messagebox.showinfo("提示","主人，活儿干完啦！")
    except Exception as e:
        if DEBUG:
            raise
        else:
            tkinter.messagebox.showinfo("错误",str(e))
    finally:
        gui.bgn_button['state'] = 'normal'

def write_report(document):

    global original_papers_lst, jcr_paragraph_position
    ei_paper_num = 0
    sci_paper_num = 0
    ei_sci_paper_num = 0
    ziyin_num = 0
    tayin_num = 0

    for paper in original_papers_lst:
        if paper['shoulu'] == 'EI':
            ei_paper_num += 1
        if paper['shoulu'] == 'SCIE':
            sci_paper_num += 1
        if paper['shoulu'] =='SCIE, EI':
            ei_sci_paper_num += 1

        if paper.get('ziyin','000') !='000':
            ziyin_num += int(paper['ziyin'])
        if paper.get('tayin','000') != '000':
            tayin_num += int(paper['tayin'])

    style = document.styles.add_style("indent", WD_STYLE_TYPE.PARAGRAPH)
    style.paragraph_format.left_indent = docx.shared.Cm(0.5)
    style.paragraph_format.space_after = docx.shared.Pt(1)
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = docx.shared.Pt(12)

    p = document.add_paragraph("")
    run = p.add_run("委托检索信息")
    run.bold = True
    run.font.size = docx.shared.Pt(14)

    document.add_paragraph("委托人：",style="indent")
    document.add_paragraph("委托人单位：", style="indent")
    document.add_paragraph("委托查询时间范围：", style="indent")

    p = document.add_paragraph("")
    run = p.add_run("检索数据库")
    run.bold = True
    run.font.size = docx.shared.Pt(14)

    i = 1
    if(len(gui.wos_input.get())>1):
        document.add_paragraph("%d. 美国《科学引文索引》（SCI-EXPANDED，网络版）" % i, style="indent")
        i += 1
    if(gui.yinyong_opt.get()==True):
        document.add_paragraph("%d. 美国《Web of Science 核心合集：引文索引》（网络版）" % i, style="indent")
        i += 1

    if (gui.jcr_opt.get() == True):
        document.add_paragraph("%d. 美国《期刊引用报告》（JCR，网络版）" % i, style="indent")
        i += 1

    if (gui.fenqu_opt.get() == True):
        document.add_paragraph("%d. 中科院SCI分区表（网络版）" % i, style="indent")
        i += 1
    if (len(gui.ei_file_path.get()) > 1):
        document.add_paragraph("%d. 美国《工程索引》(Ei Compendex，网络版）" % i, style="indent")
        i += 1

    p = document.add_paragraph("")
    run = p.add_run("检索结果")
    run.bold = True
    run.font.size = docx.shared.Pt(14)

    i = 1
    if sci_paper_num !=0:
        p = document.add_paragraph("%d.  美国《科学引文索引》(SCI-EXPANDED，网络版)收录" % i)
        i += 1
        p.add_run(" %d " % sci_paper_num).underline = True
        p.add_run("篇")
        if ziyin_num or tayin_num:
            p.add_run("，在《Web of Science 核心合集：引文索引》中累计被引用")
            p.add_run(" %d " % (ziyin_num+tayin_num)).underline = True
            p.add_run("次（其中他人引用")
            p.add_run(" %d " % tayin_num).underline = True
            p.add_run("次，自引")
            p.add_run(" %d " % ziyin_num).underline = True
            p.add_run("次。注：关于他引和自引的区分，本证明所采用的方法是：文献被除第一作者及合作者以外其他人的引用为他引）。")
    if ei_paper_num !=0:
        p = document.add_paragraph("%d. 美国《工程索引》(Ei Compendex，网络版)收录" %i)
        p.add_run(" %d " % ei_paper_num).underline = True
        p.add_run("篇")

        i += 1
    if ei_sci_paper_num !=0:
        p = document.add_paragraph("%d. EI、SCI共同收录" % i)
        p.add_run(" %d " % ei_sci_paper_num).underline = True
        p.add_run("篇")
        i += 1
    if gui.jcr_opt.get()==True:

        p = document.add_paragraph("%d. JCR收录期刊种" %i)
        jcr_paragraph_position = len(document.paragraphs)-1


    document.add_paragraph("")
    document.add_paragraph("（检索结果详见附件）")
    document.add_paragraph("特此证明！")

    document.add_paragraph("")
    p = document.add_paragraph("                        检索人：")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph("查证单位：教育部科技查新工作站（L27）")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p = document.add_paragraph("北京理工大学查新检索咨询中心")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


    p = document.add_paragraph(time.strftime("%Y{y}%m{m}%d{d}").format(y='年',m='月',d='日'))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def ei_shoulu():
    '''将EI收录的论文添加到original_paper_lst'''
    global original_papers_lst
    ei_file = gui.ei_file_path.get()
    ei_file = ei_file.strip()
    ei_papers = pd.read_excel(ei_file,0)
    i = 0
    while i<ei_papers.shape[0]:
        title = ei_papers['Title'][i]
        ei_paper = {}
        j = 0
        while j<len(original_papers_lst):
            if title == original_papers_lst[j]['title']:
                original_papers_lst[j]['accession number'] =  ei_papers['Accession number'][i]
                original_papers_lst[j]['shoulu'] += ", EI"
                break
            j += 1
        else:
            ei_paper['title'] = title
            ei_paper['author'] = ei_papers['Author'][i]
            ei_paper['source'] = ei_papers['Source'][i]
            if ei_papers.get("Volume") is not None:
                ei_paper['source'] += " 卷:"+str(ei_papers['Volume'][i]).replace("nan","")
            if ei_papers.get("Pages") is not None:
                ei_paper['source'] += "     页:"+str(ei_papers['Pages'][i]).replace("nan","")

            ei_paper['source'] += " 出版年:"+str(ei_papers['Publication year'][i])+" 文献类型:" + ei_papers['Document type'][i]

            if ei_papers.get("ISSN") is not None:
                ei_paper['issn'] = ei_papers['ISSN'][i]
            else:
                ei_paper['issn'] = ""
            ei_paper['accession number'] = ei_papers['Accession number'][i]
            ei_paper['shoulu'] = "EI"
            if ei_papers.get('Corresponding author') is not None:
                ei_paper['reprint_author'] = ei_papers['Corresponding author'][i]
            else:
                ei_paper['reprint_author'] = ""
            if ei_papers.get('Author affiliation') is not None:
                ei_paper['address'] = ei_papers['Author affiliation'][i]
            else:
                ei_paper['address'] = ""
            original_papers_lst.append(ei_paper)

        i += 1

def scrape_jcr(document):

    global original_papers_lst, jcr_paragraph_position
    jcr_shoulu_num = 0
    document.add_page_break()
    # 写入文件头
    p = document.add_paragraph("")
    p.add_run("附件二：").bold = True
    p = document.add_paragraph("")
    run = p.add_run("美国《期刊引用报告》（JCR）收录情况")
    run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
    run.bold = True

    #期刊的不重复列表
    issn_lst = []
    for paper in original_papers_lst:
        if paper['issn'] not in issn_lst:
            issn_lst.append(paper['issn'])
    try:
        db = pymysql.connect(host="118.31.57.201", db="chaxinbu", user="db_kiaora", password="Kiaora477(")
        cursor = db.cursor()
        i = 1
        for issn in issn_lst:
            cursor.execute("select * from jcr where issn='%s'" % issn)

            result = cursor.fetchone()
            if result is not None:
                title = result[1]
                jif = result[2]
                document.add_paragraph("%d. 刊名：%s" %(i,title))
                document.add_paragraph("ISSN: %s" % issn)
                document.add_paragraph("2017年影响因子: %s" % jif)
                i += 1
                jcr_shoulu_num += 1

        p = document.paragraphs[jcr_paragraph_position]
        p.text = str(p.text).replace("种","%d 种"% jcr_shoulu_num)
    except Exception as e:
        print("JCR收录出错:%s" % str(e))
    finally:
        document.save(str(gui.path_input.get()).strip() + "\\" + str(gui.bianhao_input.get()).strip() + "_baogao.docx")

def scrape_fenqu(document):
    global original_papers_lst

    document.add_page_break()
    # 写入文件头
    p = document.add_paragraph("")
    run = p.add_run("中科院SCI分区表（网络版）收录情况")
    run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
    run.bold = True

    # 期刊的不重复列表
    issn_lst = []
    for paper in original_papers_lst:
        if paper['issn'] not in issn_lst:
            issn_lst.append(paper['issn'])

    i = 1
    for issn in issn_lst:

        document.add_paragraph("%d. 刊名：" % i)
        document.add_paragraph("  ISSN：%s" % issn)
        document.add_paragraph("  2017版分区情况")

        table = document.add_table(rows=3, cols=4)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = ""
        hdr_cells[1].text = "学科名称"
        hdr_cells[2].text = "分区"
        hdr_cells[3].text = "TOP期刊"


        table.rows[1].cells[0].text = "小类"
        table.rows[2].cells[0].text = "大类"

        i += 1
    document.save(str(gui.path_input.get()).strip() + "\\" + str(gui.bianhao_input.get()).strip() + "_baogao.docx")

def report_overview(document):
    '''查询作者是否第一作者、通讯作者'''
    global original_papers_lst

    wos_jiansuo = ""  #生成WOS的 入藏号检索式，保存下来，便于重复使用
    ei_jiansuo = ""
    if len(document.paragraphs)>1:
        del document.paragraphs[:]
    author = gui.author_input.get()
    name_lst = author.split(",")
    if len(name_lst)<1:
        tkinter.messagebox.showinfo("提示","作者格式有误")
        return
    else:
        document.add_page_break()
        p = document.add_paragraph("")
        run = p.add_run("作者论文情况概览")
        run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
        run.bold = True

        #初始表格3列，序号、标题、收录
        cols = 3

        if gui.yinyong_opt.get()==True:
            cols += 1

        if gui.contribution_opt.get()==True:
            cols += 1

        table = document.add_table(rows=len(original_papers_lst)+1, cols=cols)
        table.style = "Table Grid"
        table.autofit = True

        icols = 0
        #表格包括1.序号，2.文章题目 3.收录情况 4,引用情况  5.贡献情况
        hdr_cells = table.rows[icols+0].cells
        hdr_cells[icols+0].text = "序号"
        hdr_cells[icols+1].text = "标题/作者/来源"
        hdr_cells[icols+1].width = 5486400
        icols += 1

        hdr_cells[icols+1].text = "收录情况"
        icols += 1
        if gui.yinyong_opt.get() == True:
            hdr_cells[icols+1].text = "引用情况"
            icols += 1
        if gui.contribution_opt.get()==True:
            hdr_cells[icols+1].text = "贡献情况"

        capital_pattern = re.compile("[A-Z]")

        capital_lst = capital_pattern.findall(str(name_lst[1]))
        capital_name = ""
        for capital in capital_lst:
            capital_name += capital+"[\.-]?"
        #生成正则"Mao[\s,]+E[\.-]?K[\.-]?"
        pattern_str = str(name_lst[0]).strip() + "[\s,]+" + capital_name
        paper_num = 1
        for paper in original_papers_lst:

            icols = 0
            #匹配是否第一作者
            matchObj = re.match(pattern_str,paper['author'], re.M)
            if matchObj:
                paper['bool_first_author'] = True
            else:
                #这一部分主要是针对EI文献的作者形式来测试，EI中都是作者全称
                #生成 Mao[\s,]+ErKe|ErKe[\s,]+Mao
                ptn_str = name_lst[0].strip()+"[\s,]+"+name_lst[1].strip() + "|" + name_lst[1].strip() + "[\s,]+" + name_lst[0].strip()
                #将作者中 Mao, Er-Ke调整为 Mao, ErKe进行测试
                author = paper['author'].replace("-","")
                matchObj = re.match(ptn_str,author,re.I | re.M)
                if matchObj:
                    paper['bool_first_author'] = True
                else:
                    paper['bool_first_author'] = False

            #匹配是否通讯作者
            reprint_author = str(paper['reprint_author'])
            if len(reprint_author)>1:
                matchObj = re.search(pattern_str,reprint_author,re.M)
                if matchObj:
                    paper['bool_reprint_author'] = True
                else:
                    # 这一部分主要是针对EI文献的作者形式来测试，EI中都是作者全称
                    # 生成 Mao[\s,]+ErKe|ErKe[\s,]+Mao
                    ptn_str = name_lst[0].strip() + "[\s,]+" + name_lst[1].strip() + "|" + name_lst[1].strip() + "[\s,]+" + \
                              name_lst[0].strip()
                    # 将作者中 Mao, Er-Ke调整为 Mao, ErKe进行测试
                    author = reprint_author.replace("-", "")
                    matchObj = re.search(ptn_str, author, re.I | re.M)
                    if matchObj:
                        paper['bool_reprint_author'] = True
                    else:
                        paper['bool_reprint_author'] = False
            else:
                paper['bool_reprint_author'] = False

            paper_cells = table.rows[paper_num].cells
            paper_cells[icols].text = str(paper_num)
            icols += 1
            paper_cells[icols].text = paper['title']


            if paper.get("full_author","not") != "not":
                paper_cells[icols].text += "\n 作者:" + str(paper['full_author'])
            else:
                paper_cells[icols].text += "\n 作者:" + str(paper['author'])

            paper_cells[icols].text += "\n 来源:" + paper['source']

            if paper.get("wos_no", 'not') != 'not':
                paper_cells[icols].text += "\n" + str(paper['wos_no'])
                if wos_jiansuo == "":
                    wos_jiansuo += "UT="+ str(paper['wos_no'])
                else:
                    wos_jiansuo += " OR UT="+str(paper['wos_no'])
            if paper.get('accession number', 'not') != 'not':
                paper_cells[icols].text += "\n Accession Number:" + str(paper['accession number'])
                if ei_jiansuo =="":
                    ei_jiansuo += str(paper['accession number']) + " WN AN"
                else:
                    ei_jiansuo +=" OR " + str(paper['accession number']) + " WN AN"

            icols += 1
            paper_cells[icols].text = paper['shoulu']
            if paper.get("hc", "not") != "not":
                hc = str(paper['hc']).replace("nan","")
                if len(hc)>1:
                    paper_cells[icols].text += ", 高被引论文"
            icols += 1
            if(gui.yinyong_opt.get()==True):
                if paper.get('ziyin','000') != '000':
                    paper_cells[icols].text = "自引%d次，他引%d次" %(paper['ziyin'],paper['tayin'])

                icols += 1
            if gui.contribution_opt.get()==True:
                paper_cells[icols].text = ""
                if paper['bool_first_author']:
                    paper_cells[icols].text += "第一作者  "
                if paper['bool_reprint_author']:
                    paper_cells[icols].text +="通讯作者"
            paper_num += 1

        if wos_jiansuo != "":
            f = open(str(gui.path_input.get()).strip() + "\\"+ "wos_jiansuo.txt", "w")
            f.write(wos_jiansuo)
            f.close()
        if ei_jiansuo != "":
            f = open(str(gui.path_input.get()).strip() + "\\"+"ei_jiansuo.txt","w")
            f.write(ei_jiansuo)
            f.close()
        document.save(str(gui.path_input.get()).strip() + "\\" + str(gui.bianhao_input.get()).strip() + "_baogao.docx")



class Spider_gui(object):

    def select_path(self):
        rpt_dir = tkinter.filedialog.askdirectory()
        self.path.set(rpt_dir)

    def select_ei_path(self):
        ei_file = tkinter.filedialog.askopenfilename()
        self.ei_file_path.set(ei_file)

    def select_wos_path(self):
        wos_file = tkinter.filedialog.askopenfilename()
        self.wos_file_path.set(wos_file)

    def __init__(self):
        self.window = tkinter.Tk()

        self.path = tkinter.StringVar()
        self.jcr_opt = tkinter.BooleanVar()
        self.fenqu_opt = tkinter.BooleanVar()
        self.yinyong_opt = tkinter.BooleanVar()
        self.contribution_opt = tkinter.BooleanVar()
        self.hcp_opt = tkinter.BooleanVar()

        self.progress_value = tkinter.IntVar()
        self.ei_file_path = tkinter.StringVar()
        self.wos_file_path = tkinter.StringVar()

        self.jcr_opt.set(False)
        self.yinyong_opt.set(False)
        self.fenqu_opt.set(False)

        self.window.title("检索报告 by 北理工图书馆 不懂如山")
#        self.window.iconbitmap("working.ico")

        self.wos_label = tkinter.Label(self.window, text="WOS文件:")
        self.wos_input = tkinter.Entry(self.window, width=50, textvariable=self.wos_file_path)
        self.wos_path_button = tkinter.Button(self.window, text="选择文件", command=self.select_wos_path)

        self.path_label = tkinter.Label(self.window, text="保存路径：")
        self.path_input = tkinter.Entry(self.window, width=50, textvariable=self.path)
        self.path_button = tkinter.Button(self.window, text="路径选择", command=self.select_path)

        self.ei_path_label = tkinter.Label(self.window, text="EI文件：")
        self.ei_path_input = tkinter.Entry(self.window, width=50, textvariable=self.ei_file_path)
        self.ei_path_button = tkinter.Button(self.window, text="选择文件", command=self.select_ei_path)

        self.bianhao_label = tkinter.Label(self.window, text="编号")
        self.bianhao_input = tkinter.Entry(self.window, width=10)

        self.JCR_checkbutton = tkinter.Checkbutton(self.window, text="JCR", onvalue=True, offvalue=False, width=15,
                                                   variable=self.jcr_opt)
        self.fenqu_checkbutton = tkinter.Checkbutton(self.window, text="中科院分区", onvalue=True, offvalue=False,
                                                     width=15, variable=self.fenqu_opt)
        self.contribution_checkbutton = tkinter.Checkbutton(self.window,text="贡献", onvalue=True, offvalue=False,width=15,variable=self.contribution_opt)
        self.hcp_checkbutton = tkinter.Checkbutton(self.window,text="高被引",onvalue=True, offvalue=False, width=15, variable=self.hcp_opt)
        self.yinyong_checkbutton = tkinter.Checkbutton(self.window,text="引用", onvalue=True, offvalue=False, width=15,variable=self.yinyong_opt)

        self.author_label = tkinter.Label(self.window,text="委托人英文名：")
        self.author_input = tkinter.Entry(self.window,width=50)
        self.author_tip = tkinter.Label(self.window,text="(示例: Mao, ErKe)")


        self.progress_bar = tkinter.ttk.Progressbar(self.window,orient="horizontal", length=350, mode='determinate',variable=self.progress_value, maximum=100)
        self.processing_info = tkinter.Listbox(self.window, width=50)

        self.bgn_button = tkinter.Button(self.window, command=self.begin_crawl, text="开始")

    def gui_arrange(self):
        self.wos_label.grid(row=1, column=1)
        self.wos_input.grid(row=1, column=2)
        self.wos_path_button.grid(row=1, column=3)

        self.path_label.grid(row=2, column=1)
        self.path_input.grid(row=2, column=2)
        self.path_button.grid(row=2, column=3)

        self.ei_path_label.grid(row=3, column=1)
        self.ei_path_input.grid(row=3, column=2)
        self.ei_path_button.grid(row=3, column=3)

        self.bianhao_label.grid(row=4, column=1)
        self.bianhao_input.grid(row=4, column=2, sticky="w")
        self.JCR_checkbutton.grid(row=4, column=2 )
        self.fenqu_checkbutton.grid(row=4, column=2, sticky='e')


        self.yinyong_checkbutton.grid(row=5, column=2, sticky="w")
        self.contribution_checkbutton.grid(row=5,column=2)
        self.hcp_checkbutton.grid(row=5,column=2, padx=15,sticky="e")

        self.author_label.grid(row=6, column=1)
        self.author_input.grid(row=6,column=2, stick="w")
        self.author_tip.grid(row=6,column=3,stick="w")

        self.progress_bar.grid(row=7,column=2)
        self.processing_info.grid(row=8, column=2)
        self.bgn_button.grid(row=9, column=2, sticky="e")

    def begin_crawl(self):

        url = str(self.wos_input.get()).strip()
        save_path = str(self.path.get()).strip()
        rpt_num = str(self.bianhao_input.get()).strip()
        author = str(self.author_input.get()).strip()
        try:
            urllib.parse.urlparse(url)
            if(len(rpt_num)<1):
                tkinter.messagebox.showinfo("提示","报告编号不能为空")
                return
            if(len(save_path)<1):
                tkinter.messagebox.showinfo("提示", "保存路径不能为空")
                return
            if (len(author) < 1):
                tkinter.messagebox.showinfo("提示", "作者不能为空")
                return
        except Exception as e:
            tkinter.messagebox.showinfo("错误",str(e))
            return


        t1 = threading.Thread(target=scrape_sci,args=(url,))
        t1.start()
        gui.bgn_button['state'] = 'disabled'


if __name__ == '__main__':
    gui = Spider_gui()
    gui.gui_arrange()
    # 主程序执行
    tkinter.mainloop()
